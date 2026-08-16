import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("MANUAL PENGGUNA KESELURUHAN SISTEM\nDARE TO CHANGE (D2C)")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(15, 23, 42) # Deep Navy

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Sekolah Menengah Kebangsaan Sungai Damit, Tamparuli, Sabah\nSistem Pengurusan Kehadiran, Sahsiah, Merit, Disiplin & Kaunseling Digital")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Callout Banner Box
    table_box = doc.add_table(rows=1, cols=1)
    table_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table_box.cell(0, 0)
    set_cell_background(cell, "F0F9FF") # Light Blue Tint
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_box = cell.paragraphs[0]
    r_box = p_box.add_run("📌 TEMA PROGRAM D2C: ")
    r_box.bold = True
    r_box.font.color.rgb = RGBColor(2, 132, 199)
    
    r_box2 = p_box.add_run("\"Hadir Hari Ini, Menang Esok Hari\"\n")
    r_box2.bold = True
    r_box2.font.color.rgb = RGBColor(15, 23, 42)
    
    r_box3 = p_box.add_run("⚡ TAGLINE: ")
    r_box3.bold = True
    r_box3.font.color.rgb = RGBColor(217, 119, 6)
    
    r_box4 = p_box.add_run("\"Saya Hadir, Saya Kekal, Saya Berjaya!\"\n")
    r_box4.bold = True
    r_box4.font.color.rgb = RGBColor(15, 23, 42)
    
    r_box5 = p_box.add_run("🏆 MOTTO SEKOLAH: ")
    r_box5.bold = True
    r_box5.font.color.rgb = RGBColor(16, 185, 129)
    
    r_box6 = p_box.add_run("\"ONE TEAM ONE DREAM, TERUS MARA MENAWAN 7SUMMITs\"")
    r_box6.bold = True
    r_box6.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Helper function to add section
    def add_section(num, title_text, items):
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(14)
        p_sec.paragraph_format.space_after = Pt(6)
        r = p_sec.add_run(f"{num}. {title_text}")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(15, 23, 42)

        for sub_num, sub_title, bullets in items:
            p_subhead = doc.add_paragraph()
            p_subhead.paragraph_format.space_before = Pt(6)
            p_subhead.paragraph_format.space_after = Pt(2)
            r_subhead = p_subhead.add_run(f"  {sub_num} {sub_title}")
            r_subhead.bold = True
            r_subhead.font.size = Pt(11)
            r_subhead.font.color.rgb = RGBColor(30, 41, 59)

            for b in bullets:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_before = Pt(0)
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.paragraph_format.left_indent = Inches(0.4)
                
                # Split first bold part if formatted as "Bold Part: Rest"
                if ":" in b:
                    parts = b.split(":", 1)
                    r_b1 = p_bullet.add_run(parts[0] + ":")
                    r_b1.bold = True
                    r_b1.font.color.rgb = RGBColor(15, 23, 42)
                    r_b2 = p_bullet.add_run(parts[1])
                    r_b2.font.color.rgb = RGBColor(51, 65, 85)
                else:
                    r_b = p_bullet.add_run(b)
                    r_b.font.color.rgb = RGBColor(51, 65, 85)

    # Section 1: Laman Utama Public & Landing Page
    add_section("1", "🌐 Laman Utama Public & Landing Page (https://d2csummit.online/)", [
        ("1.1", "Pusat Perhatian (Spotlight D2C)", [
            "Rekabentuk High-Tech: Menggunakan tema warna 'Deep Space Navy' dan 'Cosmic Violet' dengan kesan kad 'Glassmorphism' yang futuristik.",
            "Lampu Uptime Status: Memaparkan indikator hijau '🟢 TAMPARULI • D2C SYSTEM ONLINE' di bahagian atas header.",
            "Sasaran Program: Dilaksanakan secara menyeluruh untuk semua murid SMK Sungai Damit merangkumi Tingkatan 1, 2, 3, 4 & 5 (Sesi Pagi & Sesi Petang).",
            "Motto & Tagline: Menonjolkan slogan utama sekolah 'ONE TEAM ONE DREAM, TERUS MARA MENAWAN 7SUMMITs' dan 'Saya Hadir, Saya Kekal, Saya Berjaya!'.",
        ]),
        ("1.2", "Model 3 Aras Intervensi D2C", [
            "Aras 1 (Universal): Untuk semua murid (T1 - T5) merangkumi imbasan rekod harian, cabaran kelas, pemantauan masa rehat, dan merit sahsiah.",
            "Aras 2 (Bersasar): Untuk murid lewat berulang / berisiko, merangkumi check-in mentor UBK/PRS, pelan 5 hari 'Saya Kembali Hari Ini', dan makluman peribadi penjaga.",
            "Aras 3 (Intensif): Kes berisiko tinggi merangkumi Program Ziarah Cakna ke rumah, perjumpaan khas Pengetua/PK HEM, serta pelan pemulihan UBK & Disiplin.",
        ]),
        ("1.3", "Pautan Akses Portal (Launchpad)", [
            "Peti Pautan Selamat: Memaparkan 3 butang pintar untuk laluan Guru, Ibu Bapa, dan Murid secara tidak terlalu keterlaluan (tidak overshadow pengurusan sekolah).",
            "Portal Guru & Pentadbir: Mengarahkan ke laluan log masuk staf (/#/sign-in).",
            "Portal Ibu Bapa: Mengarahkan ke laluan carian No. IC Penjaga (/#/parent).",
            "Portal Murid: Mengarahkan ke laluan imbasan Kod QR Name Tag (/#/student).",
        ])
    ])

    # Section 2: Log Masuk & Kawalan Akses
    add_section("2", "🔑 Log Masuk & Kawalan Akses Staf (Auth & RBAC)", [
        ("2.1", "Log Masuk Pengurusan Staf", [
            "E-mel & Kata Laluan: Guru dan pentadbir sekolah log masuk melalui Supabase Auth.",
            "Semakan Konfigurasi: Sistem menyemak e-mel yang berdaftar dalam profil staf sekolah secara automatik.",
        ]),
        ("2.2", "Kawalan Hak Akses (Role-Based Access Control - RBAC)", [
            "Peranan Admin (Pentadbir): Akses penuh ke semua modul, tetapan sistem, pengurusan akaun guru, dan laporan.",
            "Peranan Disiplin (Guru Disiplin): Akses khas ke Modul Disiplin SSDOP, pendaftaran kes, dan penerbitan Pengumuman Disiplin.",
            "Peranan Kaunselor (Guru UBK): Akses khas ke Modul Kaunseling UBK, rekod sesi, Peti Suara Murid, dan penerbitan Pengumuman Kaunseling.",
            "Peranan Guru Kelas / Subjek: Akses ke imbasan QR harian, penandaan manual kelas, dan paparan merit murid.",
        ])
    ])

    # Section 3: Modul Imbasan QR & Kehadiran Harian
    add_section("3", "📷 Modul Imbasan QR & Kehadiran Harian", [
        ("3.1", "Imbasan Kad QR Name Tag Murid", [
            "Pencegahan Impersonasi: Murid diimbas menggunakan Kod QR unik pada Kad Name Tag murid (bukan No. IC) untuk mengelakkan penipuan kehadiran.",
            "Imbasan Nisbah Sesi: Waktu cutoff automatik mengikut Sesi Pagi (Tingkatan 3, 4, 5) dan Sesi Petang (Tingkatan 1, 2) serta hari persekolahan.",
            "Derivasi Status Otomatik: Imbasan sebelum cutoff dikira 'Hadir', imbasan selepas cutoff dikira 'Lewat'.",
        ]),
        ("3.2", "Kemas Kemaskini Manual & Kemaskini Pukal", [
            "Kemaskini Eksepsi Pukal: Guru boleh menandakan kelas dengan 'Default Hadir' dan hanya mendaftarkan murid yang 'Tidak Hadir', 'Cuti Sakit (MC)', atau 'Urusan Rasmi'.",
            "Kebenaran Kemaskini: Rekod dikunci dengan jejak audit (audit_log) untuk mengelakkan pemalsuan rekod.",
        ]),
        ("3.3", "Pendaftaran & Gantian Kad QR Murid", [
            "Pendaftaran Guru: Mana-mana guru boleh mendaftarkan atau menggantikan Kad QR murid yang hilang/rosak tanpa menunggu kebenaran Admin.",
        ])
    ])

    # Section 4: Modul Mata Merit & Pengiktirafan Sahsiah
    add_section("4", "🏆 Modul Mata Merit & Pengiktirafan Sahsiah", [
        ("4.1", "Sistem 4 Mata Merit Harian D2C", [
            "Rutin Step 01 (Hadir Ke Sekolah): +1 Mata Merit apabila hadir pada hari persekolahan.",
            "Rutin Step 02 (Tepat Masa): +1 Mata Merit apabila berada di kelas pada waktu ditetapkan.",
            "Rutin Step 03 (Kembali Selepas Rehat): +1 Mata Merit apabila masuk kelas selepas waktu rehat tanpa berkeliaran.",
            "Rutin Step 04 (Kekal Tamat Sesi): +1 Mata Merit apabila mengikuti PdP sehingga tamat sesi persekolahan.",
        ]),
        ("4.2", "Mata Merit Tambahan & Pengecualian", [
            "Bonus Sahsiah terpuji: Guru boleh menambah mata bonus bagi aktiviti khas sekolah.",
            "Penilaian Dinamik: Mata merit dijumlahkan secara langsung dalam pembolehubah tanpa disimpan secara statik.",
        ]),
        ("4.3", "Leaderboard & Sijil Pengiktirafan", [
            "Carta Kedudukan: Memaparkan kedudukan murid & kelas tertinggi mengikut kriteria mingguan/bulanan.",
            "6 Kategori Anugerah: Pengiktirafan automatik bagi kehadiran penuh dan peningkatan sahsiah.",
        ])
    ])

    # Section 5: Modul Papan Pemuka Analytics (Dashboard)
    add_section("5", "📊 Modul Papan Pemuka Analytics (Dashboard)", [
        ("5.1", "Kad Statistik Utama (Present Breakdown Alignment)", [
            "Penyelarasan Kad Hadir: Kad statistik 'Present' memaparkan jumlah keseluruhan (contoh: 328 Total Present: 320 Hadir • 7 MC • 1 Rasmi).",
            "Kad Stat Lengkap: Hadir (Green), Tidak Hadir (Red), Lewat (Amber), dan Cuti/Rasmi (Blue).",
        ]),
        ("5.2", "Graf Analytics & Visualisasi", [
            "Graf Trend Kehadiran: Graf garis mingguan/bulanan mengikut peratusan sekolah.",
            "Taburan Masa Imbasan: Graf taburan waktu ketibaan murid di pintu pagar/sekolah.",
            "Heatmap Kehadiran: Visual perbandingan kehadiran mengikut hari persekolahan.",
            "Kedudukan Kelas: Menampilkan 5 Kelas Terbaik & Toggle 5 Kelas Perlu Perhatian.",
        ]),
        ("5.3", "Penyusunan Kad Boleh Ubah (Drag & Drop / Button Fallback)", [
            "Drag to Reorder: Guru boleh menyusun semula posisi kad statistik mengikut keutamaan.",
            "Butang Up/Down Fallback: Sokongan navigasi pada tablet/iPad Safari.",
        ])
    ])

    # Section 6: Modul Disiplin & Kaunseling (SSDOP / UBK)
    add_section("6", "⚖️ Modul Disiplin & Kaunseling (SSDOP / UBK - /#/discipline-counseling)", [
        ("6.1", "Tab 1: Kes Disiplin SSDOP", [
            "Pendaftaran Kes SSDOP: Pendaftaran kes mengikut tahap Kes Ringan, Sederhana, dan Berat.",
            "Penjejakan Status: Status Kes Dalam Siasatan, Dirujuk UBK, atau Selesai.",
            "✍️ Special Announcement (Discipline): Borang pengumuman khas Guru Disiplin (Warna Coklat/Emas) untuk menerbitkan pesanan terus ke Portal Murid dan menyalin format teks WhatsApp PIBG.",
        ]),
        ("6.2", "Tab 2: Sesi Kaunseling UBK", [
            "Rekod Sesi Kaunseling: Pendaftaran sesi Kaunseling Individu, Kelompok, Kerjaya, dan Sahsiah.",
            "Tindakan Susulan: Penjejakan status Memerlukan Susulan atau Selesai.",
            "✍️ Special Announcement (Kaunseling UBK): Borang pengumuman khas Guru UBK (Warna Ungu UBK) untuk menerbitkan pengumuman terus ke Portal Murid dan menyalin format teks WhatsApp PIBG.",
        ]),
        ("6.3", "Tab 3: Peti Suara Murid (Teacher Inbox)", [
            "Inbox Suara Murid: Tempat Guru Kaunselor & Guru Disiplin membaca hantaran murid.",
            "Perlindungan Identiti (Anonymous): Bagi aduan buli/keselamatan yang dihantar secara rahsia, nama & kelas murid dilindungi sebagai 'SULIT / RAHSIA (ANONYMOUS)'.",
            "Maklum Balas Guru: Guru boleh mengemaskini status dan menulis nota respon yang akan dibaca oleh murid di portal mereka.",
        ]),
        ("6.4", "Tab 4: Ringkasan & Analisis", [
            "Analisis Statistik: Paparan pecahan kes disiplin & sesi kaunseling mengikut bulan dan kategori.",
        ])
    ])

    # Section 7: Modul Laporan & Eksport (Reports & WhatsApp Generator)
    add_section("7", "📈 Modul Laporan & Eksport WhatsApp", [
        ("7.1", "Laporan Trend KPI D2C", [
            "KPI Kadar Kehadiran: Graf perbandingan mingguan berasaskan sasaran program D2C.",
            "Laporan Murid Berisiko (At-Risk): Senarai murid dengan kadar kehadiran di bawah ambang (70%, 80%, 90%).",
            "Laporan Ketidakhadiran Berulang (Repeat Absence): Penjejakan murid tidak hadir berulang.",
            "Laporan Chronic Latecomers: Senarai murid lewat 3 kali atau lebih dalam tempoh 7 hari.",
            "Pecahan Jenis Cuti (Leave-Type Breakdown): Perbandingan Tidak Hadir vs Cuti Sakit (MC) vs Urusan Rasmi.",
        ]),
        ("7.2", "Penjana Laporan WhatsApp PIBG (WhatsApp Report Generator)", [
            "Format Mesra WhatsApp: Menjanakan teks laporan kehadiran kelas & tingkatan mengikut format rasmi sekolah.",
            "Pecahan Mengikut Sesi: Menjana format berasingan untuk Sesi Pagi dan Sesi Petang.",
            "Butang Salin 1-Klik: Memudahkan guru menyalin laporan terus ke kumpulan WhatsApp PIBG / Sekolah.",
        ])
    ])

    # Section 8: Portal Ibu Bapa (Parent Portal - /#/parent)
    add_section("8", "👨‍👩‍👧‍👦 Portal Ibu Bapa (Parent Portal - /#/parent)", [
        ("8.1", "Carian No. IC Penjaga (MyKad Lookup)", [
            "Tanpa Log Masuk Rumit: Ibu bapa hanya perlu memasukkan No. IC / MyKad Penjaga yang berdaftar.",
            "Kadar Had Limit (Rate Limiting): Dilindungi dengan kawalan keselamatan carian pangkalan data.",
        ]),
        ("8.2", "Sokongan Berbilang Anak (Multi-Sibling Tab Bar)", [
            "Tab Berbilang Anak: Jika penjaga mempunyai lebih daripada seorang anak di SMK Sungai Damit, sistem memaparkan tab untuk setiap anak secara automatik.",
        ]),
        ("8.3", "Paparan Rekod Anak", [
            "Statistik Harian: Memaparkan peratus kehadiran %, jumlah hari hadir/tidak hadir, dan jumlah mata merit anak.",
            "Status Terkini: Memaparkan status kehadiran terkini real-time anak di sekolah.",
        ])
    ])

    # Section 9: Portal Murid & Suara Murid (/#/student)
    add_section("9", "🎓 Portal Murid & Suara Murid (/#/student)", [
        ("9.1", "Log Masuk Selamat Kad QR Name Tag", [
            "Imbasan Kamera Kad QR: Murid mengimbas Kod QR pada kad nama fizikal mereka atau memasukkan 8-aksara kod token.",
            "Pencegahan Penyamaran: Log masuk disahkan melalui jadual qr_tokens pangkalan data.",
        ]),
        ("9.2", "Tab 1: 📢 Pengumuman (BAHARU)", [
            "Paparan Pengumuman Rasmi: Memaparkan Pengumuman Disiplin (Coklat/Emas) dan Pengumuman Kaunseling UBK (Ungu) yang diterbitkan oleh guru secara langsung (Live Sync).",
            "Info Murid Terlibat: Memaparkan maklumat murid yang dilampirkan secara khusus jika ada.",
        ]),
        ("9.3", "Tab 2: 📈 Kemajuan Saya", [
            "Rekod Kehadiran 30 Hari: Memaparkan cip status harian 30 hari terkini (Hadir, Lewat, Tidak Hadir).",
            "Mata Merit Terkumpul: Memaparkan lencana merit individu dan peratusan kehadiran murid.",
        ]),
        ("9.4", "Tab 3: 📣 Suara Murid (Peti Cadangan & Aduan Rahsia)", [
            "4 Kategori Hantaran: (1) Cadangan Sekolah, (2) Maklum Balas Pembelajaran, (3) Aduan Buli & Keselamatan, (4) Permohonan Kaunseling UBK.",
            "Pilihan Hantar Secara Rahsia (Anonymous): Murid boleh menyembunyikan nama & kelas bagi aduan buli demi keselamatan.",
            "Penjejak Status Respon: Murid boleh menyemak status hantaran mereka (Baru, Dalam Tindakan, Selesai) dan membaca maklum balas guru.",
        ]),
        ("9.5", "Tab 4: 🌟 Inspirasi & Pengumuman", [
            "Kata-kata Semangat: Slogan motivasi sahsiah & info kejayaan program D2C.",
        ])
    ])

    # Section 10: Tetapan Sistem & Pengurusan Akaun (Settings)
    add_section("10", "⚙️ Tetapan Sistem & Pengurusan Akaun (Settings)", [
        ("10.1", "Tetapan Tempoh Program & Merit", [
            "Pelarasan Tarikh Program: Tetapan tarikh mula dan tamat Program D2C.",
            "Kawalan Komponen Merit: Kebenaran mengaktifkan / mematikan rutin merit secara fleksibel.",
        ]),
        ("10.2", "Tetapan Masa Masuk (Cutoff Times)", [
            "Masa Cutoff Sesi: Tetapan waktu cutoff berasingan untuk Sesi Pagi dan Sesi Petang mengikut hari persekolahan (Isnin - Jumaat).",
        ]),
        ("10.3", "Pengurusan Akaun Guru & Staf", [
            "Pendaftaran Staf: Penambahan e-mel guru dan penentuan peranan (Admin, Disiplin, Kaunselor, Teacher).",
        ])
    ])

    # Save document
    output_path = os.path.join("docs", "MANUAL_PENGGUNA_D2C.docx")
    doc.save(output_path)
    print(f"Successfully generated Word Document at: {output_path}")

if __name__ == "__main__":
    generate_docx()
